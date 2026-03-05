"""
app/services/parser.py
───────────────────────
Multi-format document parser.

Supported formats: PDF, DOCX, XLSX, TXT
Returns clean plain text for each document.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import BinaryIO

from loguru import logger

from app.core.exceptions import DocumentParseError, UnsupportedFileTypeError

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".txt", ".csv"}


def extract_text(file_path: str | Path) -> str:
    """
    Extract plain text from a file.

    Args:
        file_path: Path to file on disk.

    Returns:
        Extracted text string.

    Raises:
        UnsupportedFileTypeError: If file type is not supported.
        DocumentParseError: If extraction fails.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"File type '{ext}' is not supported.",
            detail=f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}",
        )

    logger.debug(f"Extracting text from {path.name} ({ext})")

    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        elif ext == ".docx":
            return _extract_docx(path)
        elif ext in (".xlsx", ".xls"):
            return _extract_xlsx(path)
        elif ext in (".txt", ".csv"):
            return path.read_text(encoding="utf-8", errors="ignore")
    except (UnsupportedFileTypeError, DocumentParseError):
        raise
    except Exception as exc:
        raise DocumentParseError(
            f"Failed to parse '{path.name}'",
            detail=str(exc),
        ) from exc

    return ""


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """
    Extract text from raw bytes (e.g. from an UploadFile).

    Args:
        content: Raw file bytes.
        filename: Original filename (used to determine type).

    Returns:
        Extracted text string.
    """
    ext = Path(filename).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"File type '{ext}' is not supported.",
            detail=f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}",
        )

    logger.debug(f"Extracting text from bytes: {filename} ({ext})")

    try:
        if ext == ".pdf":
            return _extract_pdf_bytes(content)
        elif ext == ".docx":
            return _extract_docx_bytes(content)
        elif ext in (".xlsx", ".xls"):
            return _extract_xlsx_bytes(content)
        elif ext in (".txt", ".csv"):
            return content.decode("utf-8", errors="ignore")
    except (UnsupportedFileTypeError, DocumentParseError):
        raise
    except Exception as exc:
        raise DocumentParseError(
            f"Failed to parse '{filename}'",
            detail=str(exc),
        ) from exc

    return ""


# ── PDF ───────────────────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i + 1}]\n{text.strip()}")
    return "\n\n".join(pages)


def _extract_pdf_bytes(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i + 1}]\n{text.strip()}")
    return "\n\n".join(pages)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Tables (flatten to readable text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _extract_docx_bytes(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


# ── XLSX ──────────────────────────────────────────────────────────────────────

def _extract_xlsx(path: Path) -> str:
    import pandas as pd
    xl = pd.ExcelFile(str(path))
    sheets: list[str] = []
    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name, header=None)
        df = df.dropna(how="all")
        text = df.to_string(index=False, header=False, na_rep="")
        if text.strip():
            sheets.append(f"[Sheet: {sheet_name}]\n{text.strip()}")
    return "\n\n".join(sheets)


def _extract_xlsx_bytes(content: bytes) -> str:
    import pandas as pd
    xl = pd.ExcelFile(io.BytesIO(content))
    sheets: list[str] = []
    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name, header=None)
        df = df.dropna(how="all")
        text = df.to_string(index=False, header=False, na_rep="")
        if text.strip():
            sheets.append(f"[Sheet: {sheet_name}]\n{text.strip()}")
    return "\n\n".join(sheets)


# ── Utility ───────────────────────────────────────────────────────────────────

def truncate_text(text: str, max_chars: int = 12000) -> str:
    """
    Truncate text to max_chars to stay within LLM context limits.
    Preserves beginning and end with a [TRUNCATED] marker in the middle.
    """
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return (
        text[:half]
        + f"\n\n[... TRUNCATED — {len(text) - max_chars} chars omitted ...]\n\n"
        + text[-half:]
    )