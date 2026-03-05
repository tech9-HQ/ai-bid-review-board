"""
app/services/output_generator.py
──────────────────────────────────
Stage 6: Generate all output documents.

Produces:
  • Board_Scorecard.docx   — Governance scorecard + issue log
  • Issue_Log.xlsx         — Machine-readable findings register
  • Proposal_Revised_v1.docx  — Corrected proposal
  • SOW_Revised_v1.docx    — Legal-reviewed SOW clauses
"""
from __future__ import annotations

import uuid
from datetime import datetime
from pathlib import Path

from loguru import logger

from app.core.config import settings
from app.core.exceptions import OutputGenerationError
from app.models.schemas import ReviewSession, Severity, Verdict


def generate_all_outputs(session: ReviewSession) -> dict[str, str]:
    """
    Generate all 4 output files.

    Returns:
        Dict of {document_role: file_path_string}
    """
    settings.ensure_dirs()
    session_dir = settings.OUTPUT_DIR / session.session_id
    session_dir.mkdir(parents=True, exist_ok=True)

    output_files: dict[str, str] = {}

    try:
        output_files["scorecard"] = _generate_scorecard_docx(session, session_dir)
        output_files["issue_log"] = _generate_issue_log_xlsx(session, session_dir)
        output_files["proposal"] = _generate_proposal_docx(session, session_dir)
        output_files["sow"] = _generate_sow_docx(session, session_dir)
    except Exception as exc:
        raise OutputGenerationError(
            "Failed to generate output documents.",
            detail=str(exc),
        ) from exc

    return output_files


# ── Board Scorecard DOCX ──────────────────────────────────────────────────────

def _generate_scorecard_docx(session: ReviewSession, output_dir: Path) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import docx.oxml

    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Cover ──────────────────────────────────────────────────────────────────
    title = doc.add_heading("AI Bid Review Board", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Governance Scorecard")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)

    doc.add_paragraph(f"Deal: {session.deal_name}")
    doc.add_paragraph(f"Session ID: {session.session_id}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph()

    # ── Summary ────────────────────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    audit = session.audit_result

    if audit:
        sc = audit.scorecard
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = "Table Grid"
        _cell(summary_table, 0, 0, "Overall Score", bold=True)
        _cell(summary_table, 0, 1, f"{sc.overall_score} / 100")
        _cell(summary_table, 1, 0, "Recommendation", bold=True)
        _cell(summary_table, 1, 1, sc.recommendation.value)
        _cell(summary_table, 2, 0, "Blockers", bold=True)
        _cell(summary_table, 2, 1, str(sc.blocker_count))
        _cell(summary_table, 3, 0, "Total Issues", bold=True)
        _cell(summary_table, 3, 1, str(len(audit.issues)))
        doc.add_paragraph()

    # ── Scorecard Table ────────────────────────────────────────────────────────
    doc.add_heading("Governance Scorecard (10 Areas)", level=1)
    if audit and audit.scorecard.areas:
        cols = ["Audit Area", "Score", "Issues", "Verdict", "Notes"]
        sc_table = doc.add_table(rows=1 + len(audit.scorecard.areas), cols=len(cols))
        sc_table.style = "Table Grid"

        # Header
        for j, col in enumerate(cols):
            _cell(sc_table, 0, j, col, bold=True)

        for i, area in enumerate(audit.scorecard.areas):
            _cell(sc_table, i + 1, 0, area.area)
            _cell(sc_table, i + 1, 1, str(area.score))
            _cell(sc_table, i + 1, 2, str(area.issue_count))
            _cell(sc_table, i + 1, 3, area.verdict.value)
            _cell(sc_table, i + 1, 4, area.notes)

        doc.add_paragraph()

    # ── Issue Log ──────────────────────────────────────────────────────────────
    doc.add_heading("Issue Log", level=1)
    if audit and audit.issues:
        issue_cols = ["ID", "Category", "Severity", "Finding", "Fix", "Owner"]
        issue_table = doc.add_table(rows=1 + len(audit.issues), cols=len(issue_cols))
        issue_table.style = "Table Grid"

        for j, col in enumerate(issue_cols):
            _cell(issue_table, 0, j, col, bold=True)

        for i, issue in enumerate(audit.issues):
            _cell(issue_table, i + 1, 0, issue.id)
            _cell(issue_table, i + 1, 1, issue.category.value)
            _cell(issue_table, i + 1, 2, issue.severity.value)
            _cell(issue_table, i + 1, 3, issue.finding)
            _cell(issue_table, i + 1, 4, issue.fix)
            _cell(issue_table, i + 1, 5, issue.owner)

        doc.add_paragraph()

    # ── Clarifying Questions ───────────────────────────────────────────────────
    if audit and audit.scorecard.clarifying_questions:
        doc.add_heading("Clarifying Questions", level=1)
        for q in audit.scorecard.clarifying_questions:
            doc.add_paragraph(f"• {q}")

    path = output_dir / "Board_Scorecard.docx"
    doc.save(str(path))
    logger.info(f"[Output] Scorecard saved: {path}")
    return str(path)


# ── Issue Log XLSX ─────────────────────────────────────────────────────────────

def _generate_issue_log_xlsx(session: ReviewSession, output_dir: Path) -> str:
    import xlsxwriter

    path = output_dir / "Issue_Log.xlsx"
    workbook = xlsxwriter.Workbook(str(path))

    # Formats
    header_fmt = workbook.add_format({
        "bold": True, "bg_color": "#1f2330", "font_color": "#e8ff47",
        "border": 1, "font_size": 10, "align": "center", "valign": "vcenter",
    })
    blocker_fmt = workbook.add_format({
        "bg_color": "#3d1519", "font_color": "#ff4766",
        "border": 1, "font_size": 9, "text_wrap": True, "valign": "top",
    })
    high_fmt = workbook.add_format({
        "bg_color": "#2d2010", "font_color": "#f59e0b",
        "border": 1, "font_size": 9, "text_wrap": True, "valign": "top",
    })
    normal_fmt = workbook.add_format({
        "border": 1, "font_size": 9, "text_wrap": True, "valign": "top",
    })
    title_fmt = workbook.add_format({
        "bold": True, "font_size": 14, "font_color": "#e8ff47",
    })

    # ── Issues Sheet ───────────────────────────────────────────────────────────
    ws = workbook.add_worksheet("Issue Log")
    ws.set_tab_color("#ff4766")

    ws.write(0, 0, f"Tech9Labs AI Bid Review Board — {session.deal_name}", title_fmt)
    ws.write(1, 0, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}", normal_fmt)

    headers = ["ID", "Category", "Severity", "Finding", "Evidence", "Impact", "Fix", "Owner", "Status"]
    col_widths = [14, 16, 10, 40, 40, 30, 40, 12, 10]

    row = 3
    for col, (h, w) in enumerate(zip(headers, col_widths)):
        ws.write(row, col, h, header_fmt)
        ws.set_column(col, col, w)

    ws.set_row(row, 20)

    audit = session.audit_result
    if audit:
        for issue in audit.issues:
            row += 1
            fmt = blocker_fmt if issue.severity == Severity.BLOCKER else (
                high_fmt if issue.severity == Severity.HIGH else normal_fmt
            )
            ws.write(row, 0, issue.id, fmt)
            ws.write(row, 1, issue.category.value, fmt)
            ws.write(row, 2, issue.severity.value, fmt)
            ws.write(row, 3, issue.finding, fmt)
            ws.write(row, 4, issue.evidence, fmt)
            ws.write(row, 5, issue.impact, fmt)
            ws.write(row, 6, issue.fix, fmt)
            ws.write(row, 7, issue.owner, fmt)
            ws.write(row, 8, issue.status.value, fmt)
            ws.set_row(row, 50)

    # ── Scorecard Sheet ────────────────────────────────────────────────────────
    ws2 = workbook.add_worksheet("Scorecard")
    ws2.set_tab_color("#e8ff47")

    ws2.write(0, 0, "Governance Scorecard", title_fmt)
    ws2.write(1, 0, f"Overall Score: {audit.scorecard.overall_score if audit else 'N/A'}", normal_fmt)
    ws2.write(2, 0, f"Recommendation: {audit.scorecard.recommendation.value if audit else 'N/A'}", normal_fmt)

    sc_headers = ["Audit Area", "Score", "Issue Count", "Verdict", "Notes"]
    sc_widths = [25, 8, 12, 10, 50]
    row = 4
    for col, (h, w) in enumerate(zip(sc_headers, sc_widths)):
        ws2.write(row, col, h, header_fmt)
        ws2.set_column(col, col, w)

    if audit:
        for area in audit.scorecard.areas:
            row += 1
            colour_fmt = (
                blocker_fmt if area.verdict == Verdict.BLOCKER else
                high_fmt if area.verdict == Verdict.REVIEW else
                normal_fmt
            )
            ws2.write(row, 0, area.area, colour_fmt)
            ws2.write(row, 1, area.score, colour_fmt)
            ws2.write(row, 2, area.issue_count, colour_fmt)
            ws2.write(row, 3, area.verdict.value, colour_fmt)
            ws2.write(row, 4, area.notes, colour_fmt)

    workbook.close()
    logger.info(f"[Output] Issue log saved: {path}")
    return str(path)


# ── Proposal DOCX ─────────────────────────────────────────────────────────────

def _generate_proposal_docx(session: ReviewSession, output_dir: Path) -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    rewrite = session.proposal_rewrite

    doc.add_heading(f"Proposal — {session.deal_name}", level=0)
    doc.add_paragraph(f"Prepared by: Tech9Labs / Tech9IQ")
    doc.add_paragraph(f"Version: Revised v1 (AI Board Reviewed)")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_page_break()

    sections = [
        ("Executive Summary", rewrite.executive_summary if rewrite else ""),
        ("Solution Approach", rewrite.solution_approach if rewrite else ""),
        ("Architecture Justification", rewrite.architecture_justification if rewrite else ""),
        ("Sizing Assumptions & Headroom", rewrite.sizing_assumptions if rewrite else ""),
        ("Scope & Deliverables", rewrite.scope_and_deliverables if rewrite else ""),
        ("Milestones & Acceptance Criteria", rewrite.milestones_and_acceptance if rewrite else ""),
        ("Dependencies & Customer Responsibilities", rewrite.dependencies if rewrite else ""),
        ("Commercial Clarifications", rewrite.commercial_clarifications if rewrite else ""),
        ("Assumptions & Exclusions", rewrite.assumptions_and_exclusions if rewrite else ""),
    ]

    for heading, content in sections:
        doc.add_heading(heading, level=1)
        if content.strip():
            doc.add_paragraph(content)
        else:
            p = doc.add_paragraph("[Content pending — no data from AI stage]")
            p.runs[0].italic = True
        doc.add_paragraph()

    path = output_dir / "Proposal_Revised_v1.docx"
    doc.save(str(path))
    logger.info(f"[Output] Proposal saved: {path}")
    return str(path)


# ── SOW DOCX ──────────────────────────────────────────────────────────────────

def _generate_sow_docx(session: ReviewSession, output_dir: Path) -> str:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    legal = session.legal_review

    doc.add_heading(f"Statement of Work — {session.deal_name}", level=0)
    doc.add_paragraph("Prepared by: Tech9Labs / Tech9IQ")
    doc.add_paragraph(f"Version: Revised v1 (Legal Board Reviewed by Claude)")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_page_break()

    doc.add_heading("Legal Notice", level=2)
    note = doc.add_paragraph(
        "This Statement of Work has been reviewed and revised by the Tech9Labs AI Legal Review Board "
        "(powered by Anthropic Claude). All clauses have been reviewed for ambiguity, liability exposure, "
        "and commercial risk. This document should be reviewed by qualified legal counsel before execution."
    )
    note.runs[0].italic = True

    clauses = [
        ("SLA Clause", legal.revised_sla_clause if legal else ""),
        ("Limitation of Liability", legal.revised_liability_clause if legal else ""),
        ("Change Control Process", legal.revised_change_control if legal else ""),
        ("Acceptance Criteria", legal.revised_acceptance_criteria if legal else ""),
        ("Warranty", legal.revised_warranty_clause if legal else ""),
    ]

    for heading, content in clauses:
        doc.add_heading(heading, level=1)
        if content.strip():
            doc.add_paragraph(content)
        else:
            p = doc.add_paragraph("[Standard clause — to be populated by legal team]")
            p.runs[0].italic = True
        doc.add_paragraph()

    if legal and legal.additional_recommendations:
        doc.add_heading("Additional Legal Recommendations", level=1)
        for rec in legal.additional_recommendations:
            doc.add_paragraph(f"• {rec}")

    path = output_dir / "SOW_Revised_v1.docx"
    doc.save(str(path))
    logger.info(f"[Output] SOW saved: {path}")
    return str(path)


# ── Helper ────────────────────────────────────────────────────────────────────

def _cell(table, row: int, col: int, text: str, bold: bool = False) -> None:
    cell = table.cell(row, col)
    cell.text = text
    if bold:
        for run in cell.paragraphs[0].runs:
            run.bold = True